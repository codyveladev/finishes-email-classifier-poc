"""Attachment text extraction. Dispatches on file extension.

Phase 4a: PDF + Word (.docx). Excel (.xlsx) is Phase 4b, deferred.
Legacy .doc / .xls need LibreOffice or pywin32 and are out of scope for the POC.
"""

from pathlib import Path

MAX_CHARS_PER_FILE = 50_000  # soft per-file cap; downstream signal builder truncates further


def extract_text(path: str | None) -> str:
    if not path:
        return ""
    p = Path(path)
    if not p.exists():
        return ""

    ext = p.suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(p)
    elif ext == ".docx":
        text = _extract_docx(p)
    else:
        return ""

    if len(text) > MAX_CHARS_PER_FILE:
        text = text[:MAX_CHARS_PER_FILE]
    return text


def _extract_pdf(p: Path) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(p) as pdf:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
        if text.strip():
            return text
    except Exception:
        pass

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        return ""


def _extract_docx(p: Path) -> str:
    """Iterate paragraphs AND tables — a paragraphs-only extractor silently
    drops table content, which is where line items, cost breakdowns, and
    contract terms typically live."""
    try:
        from docx import Document
    except ImportError:
        return ""

    try:
        doc = Document(str(p))
    except Exception:
        return ""

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append("\t".join(cells))

    return "\n".join(parts)
